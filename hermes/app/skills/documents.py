from __future__ import annotations

import re
import struct
from pathlib import Path

ALLOWED_EXTENSIONS = frozenset({".pdf", ".doc", ".docx", ".md", ".markdown", ".txt"})
MAX_FILE_SIZE = 50 * 1024 * 1024

_OLE_SIGNATURE = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
_DOC_CONTROL_MAP = {
    "\r": "\n",
    "\x07": "\n",
    "\x0b": "\n",
    "\x0c": "\n",
    "\x1e": "-",
    "\x1f": "",
    "\x0e": "",
}
_DOC_DROPPED_CHARS = frozenset("\x01\x02\x03\x04\x05\x06\x08\x13\x14\x15\x1a\x1b")


class ParserError(Exception):
    pass


def read_document(file_path: str | Path) -> str:
    """Проверяет формат и размер файла и возвращает извлеченный текст."""
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ParserError(f"Неподдерживаемый формат {ext}. Допустимы: {formats_hint()}")
    if path.stat().st_size > MAX_FILE_SIZE:
        raise ParserError("Файл превышает допустимый размер 50 МБ")
    text = extract_text(path)
    if not text.strip():
        raise ParserError("Файл не содержит текста")
    return text


def formats_hint() -> str:
    return ", ".join(sorted(ALLOWED_EXTENSIONS))


def extract_text(file_path: str | Path) -> str:
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext == ".doc":
        return _extract_doc(path)
    if ext in {".md", ".markdown", ".txt"}:
        return _extract_plain_text(path)
    raise ParserError(f"Неподдерживаемый формат файла: {ext}")


def _extract_pdf(path: Path) -> str:
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
    if not pages:
        raise ParserError("Файл содержит только изображения, текстовый слой не обнаружен")
    return "\n".join(pages)


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append("\t".join(cells))
        if rows:
            parts.append("\n".join(rows))
    if not parts:
        raise ParserError("Файл не содержит текста")
    return "\n\n".join(parts)


def _extract_plain_text(path: Path) -> str:
    for encoding in ("utf-8-sig", "cp1251"):
        try:
            text = path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
        return _normalize_lines(text)
    raise ParserError("Не удалось прочитать файл: ожидается кодировка UTF-8 или CP1251")


def _extract_doc(path: Path) -> str:
    with path.open("rb") as handle:
        head = handle.read(8)
    # .docx, переименованный в .doc, остается ZIP-контейнером
    if head[:2] == b"PK":
        return _extract_docx(path)
    if head[:5] == b"{\\rtf":
        return _extract_rtf(path)
    if head != _OLE_SIGNATURE:
        raise ParserError(
            "Файл .doc не распознан как документ Word. Пересохраните его в DOCX или PDF."
        )
    return _extract_ole_doc(path)


def _extract_ole_doc(path: Path) -> str:
    try:
        import olefile
    except ImportError as exc:  # pragma: no cover - зависимость объявлена в requirements
        raise ParserError("Для чтения .doc требуется пакет olefile") from exc

    ole = olefile.OleFileIO(str(path))
    try:
        if not ole.exists("WordDocument"):
            raise ParserError("В файле .doc нет потока WordDocument")
        document = ole.openstream("WordDocument").read()
        flags = struct.unpack_from("<H", document, 0x0A)[0]
        table_name = "1Table" if flags & 0x0200 else "0Table"
        if not ole.exists(table_name):
            raise ParserError(f"В файле .doc нет потока {table_name}")
        table = ole.openstream(table_name).read()
    finally:
        ole.close()

    fc_clx, lcb_clx = struct.unpack_from("<II", document, 0x01A2)
    clx = table[fc_clx : fc_clx + lcb_clx]
    pieces = _parse_piece_table(clx)
    parts = [_read_doc_piece(document, offset, length, compressed) for offset, length, compressed in pieces]
    text = _clean_doc_text("".join(parts))
    if not text.strip():
        raise ParserError("Файл .doc не содержит текстового слоя")
    return text


def _parse_piece_table(clx: bytes) -> list[tuple[int, int, bool]]:
    index = 0
    while index < len(clx):
        kind = clx[index]
        if kind == 0x01:  # Prc: свойства форматирования, пропускаем
            size = struct.unpack_from("<H", clx, index + 1)[0]
            index += 3 + size
            continue
        if kind == 0x02:  # Pcdt: таблица фрагментов текста
            size = struct.unpack_from("<I", clx, index + 1)[0]
            return _parse_plcpcd(clx[index + 5 : index + 5 + size])
        break
    raise ParserError("В файле .doc не найдена таблица фрагментов текста")


def _parse_plcpcd(plc: bytes) -> list[tuple[int, int, bool]]:
    count = (len(plc) - 4) // 12
    if count <= 0:
        raise ParserError("В файле .doc пустая таблица фрагментов текста")
    positions = [struct.unpack_from("<I", plc, item * 4)[0] for item in range(count + 1)]
    descriptors_at = (count + 1) * 4
    pieces: list[tuple[int, int, bool]] = []
    for item in range(count):
        fc = struct.unpack_from("<I", plc, descriptors_at + item * 8 + 2)[0]
        compressed = bool(fc & 0x40000000)
        offset = fc & 0x3FFFFFFF
        if compressed:
            offset //= 2
        pieces.append((offset, positions[item + 1] - positions[item], compressed))
    return pieces


def _read_doc_piece(document: bytes, offset: int, length: int, compressed: bool) -> str:
    if length <= 0:
        return ""
    if compressed:
        return _decode_ansi(document[offset : offset + length])
    return document[offset : offset + length * 2].decode("utf-16-le", errors="ignore")


def _decode_ansi(raw: bytes) -> str:
    """Сжатые фрагменты хранят однобайтовый текст в кодовой странице документа."""
    candidates = [raw.decode(encoding, errors="replace") for encoding in ("cp1251", "cp1252")]
    return max(candidates, key=_readability_score)


def _readability_score(text: str) -> int:
    score = 0
    for char in text:
        if char.isascii():
            score += 1
        elif "\u0400" <= char <= "\u04ff":
            score += 2
        else:
            score -= 1
    return score


def _clean_doc_text(text: str) -> str:
    chars = []
    for char in text:
        if char in _DOC_DROPPED_CHARS:
            continue
        chars.append(_DOC_CONTROL_MAP.get(char, char))
    return _normalize_lines("".join(chars))


def _extract_rtf(path: Path) -> str:
    raw = path.read_bytes().decode("cp1251", errors="ignore")
    text = re.sub(r"\\'([0-9a-fA-F]{2})", lambda match: chr(int(match.group(1), 16)), raw)
    text = re.sub(r"\\u(-?\d+)\??", lambda match: chr(int(match.group(1)) % 65536), text)
    text = re.sub(r"\\par[d]?\b", "\n", text)
    text = re.sub(r"\{\\\*.*?\}", "", text, flags=re.DOTALL)
    text = re.sub(r"\\[a-zA-Z]+-?\d*\s?", "", text)
    text = text.replace("{", "").replace("}", "")
    return _normalize_lines(text)


def _normalize_lines(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    lines = [line.rstrip() for line in text.split("\n")]
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()
