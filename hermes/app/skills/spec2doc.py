from __future__ import annotations

import logging
import re
from pathlib import Path
from urllib.parse import quote

import requests

from app import config
from app.skills.llm import generate_text
from app.skills.loader import load_instructions

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = frozenset({".pdf", ".docx"})
MAX_FILE_SIZE = 50 * 1024 * 1024

MAX_DIFF_FILES = 50
MAX_FILE_DIFF_CHARS = 4000
MAX_TOTAL_DIFF_CHARS = 20000
MAX_MR_DESCRIPTION_CHARS = 4000

_GITLAB_MR_URL_RE = re.compile(
    r"^(?P<base>https?://[^/\s]+)/(?P<project>[^\s]+?)/(?:-/)?merge_requests/(?P<iid>\d+)",
    re.IGNORECASE,
)


class ParserError(Exception):
    pass


class GitLabError(Exception):
    pass


def generate_draft_from_file(path: Path) -> str:
    ext = path.suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ParserError(f"Неподдерживаемый формат {ext}. Допустимы: .pdf, .docx")
    if path.stat().st_size > MAX_FILE_SIZE:
        raise ParserError("Файл превышает допустимый размер 50 МБ")
    text = extract_text(path)
    if not text.strip():
        raise ParserError("Файл не содержит текста постановки")
    return generate_draft(text)


def generate_draft(extracted_text: str) -> str:
    user_message = (
        "Обработай следующую постановку аналитика и сформируй черновик технической документации.\n\n"
        f"<ПОСТАНОВКА>\n{extracted_text}\n</ПОСТАНОВКА>"
    )
    logger.info("[spec2doc] model=%s chars=%d", config.LLM_MODEL_NAME, len(extracted_text))
    return generate_text(
        [
            {"role": "system", "content": load_instructions("spec2doc")},
            {"role": "user", "content": user_message},
        ],
        temperature=0.2,
        max_tokens=config.LLM_MAX_TOKENS,
    )


def generate_draft_from_merge_request(url: str, context_text: str = "") -> str:
    merge_request = fetch_merge_request(url)
    user_message = (
        "Проанализируй merge request и сформируй черновик технической документации.\n\n"
        f"<MERGE_REQUEST>\n{build_merge_request_prompt(merge_request)}\n</MERGE_REQUEST>"
    )
    comment = context_text.strip()
    if comment:
        user_message += f"\n\n<КОММЕНТАРИЙ ПОСТАНОВЩИКА>\n{comment}\n</КОММЕНТАРИЙ ПОСТАНОВЩИКА>"
    logger.info(
        "[spec2doc] merge_request=%s!%s model=%s files=%d",
        merge_request["project"],
        merge_request["iid"],
        config.LLM_MODEL_NAME,
        len(merge_request["files"]),
    )
    return generate_text(
        [
            {"role": "system", "content": load_instructions("spec2doc", "instructions-merge_request.md")},
            {"role": "user", "content": user_message},
        ],
        temperature=0.2,
        max_tokens=config.LLM_MAX_TOKENS,
    )


def parse_merge_request_url(url: str) -> tuple[str, str, int]:
    match = _GITLAB_MR_URL_RE.match(url.strip())
    if not match:
        raise GitLabError(
            "Не удалось распознать ссылку на merge request. "
            "Ожидается вид https://gitlab.example.com/group/project/-/merge_requests/42"
        )
    project = match.group("project").strip("/")
    if len([part for part in project.split("/") if part]) < 2:
        raise GitLabError("В ссылке не найден путь проекта вида group/project")
    return match.group("base").rstrip("/"), project, int(match.group("iid"))


def fetch_merge_request(url: str) -> dict:
    base_url, project, iid = parse_merge_request_url(url)
    project_id = quote(project, safe="")
    merge_request = _gitlab_get(base_url, f"projects/{project_id}/merge_requests/{iid}")
    commits = _gitlab_get(
        base_url,
        f"projects/{project_id}/merge_requests/{iid}/commits",
        params={"per_page": 100},
        allow_missing=True,
    )
    files, truncated = _collect_diffs(_get_changes(base_url, project_id, iid))
    if not files and not (merge_request.get("description") or "").strip():
        raise GitLabError("Merge request не содержит ни описания, ни изменений для анализа")
    return {
        "project": project,
        "iid": iid,
        "web_url": merge_request.get("web_url") or url.strip(),
        "title": merge_request.get("title", ""),
        "description": (merge_request.get("description") or "").strip()[:MAX_MR_DESCRIPTION_CHARS],
        "state": merge_request.get("state", ""),
        "author": (merge_request.get("author") or {}).get("name", ""),
        "source_branch": merge_request.get("source_branch", ""),
        "target_branch": merge_request.get("target_branch", ""),
        "labels": merge_request.get("labels") or [],
        "commits": [
            {"short_id": item.get("short_id", ""), "title": item.get("title", "")}
            for item in commits or []
        ],
        "files": files,
        "diff_truncated": truncated,
    }


def build_merge_request_prompt(merge_request: dict) -> str:
    header = [
        f"Merge request: !{merge_request['iid']} {merge_request.get('title', '')}".strip(),
        f"Проект: {merge_request['project']}",
        f"Ссылка: {merge_request.get('web_url', '')}",
    ]
    if merge_request.get("author"):
        header.append(f"Автор: {merge_request['author']}")
    if merge_request.get("state"):
        header.append(f"Статус: {merge_request['state']}")
    if merge_request.get("source_branch") or merge_request.get("target_branch"):
        header.append(f"Ветки: {merge_request.get('source_branch', '')} -> {merge_request.get('target_branch', '')}")
    if merge_request.get("labels"):
        header.append(f"Метки: {', '.join(merge_request['labels'])}")

    blocks = ["\n".join(header)]
    if merge_request.get("description"):
        blocks.append(f"Описание merge request:\n{merge_request['description']}")

    commits = merge_request.get("commits") or []
    if commits:
        lines = "\n".join(f"- [{item['short_id']}] {item['title']}" for item in commits)
        blocks.append(f"Коммиты ({len(commits)}):\n{lines}")

    files = merge_request.get("files") or []
    if files:
        lines = "\n".join(f"- {item['path']} ({item['status']})" for item in files)
        blocks.append(f"Измененные файлы ({len(files)}):\n{lines}")

    diffs = [f"--- {item['path']} ---\n{item['diff']}" for item in files if item["diff"]]
    if diffs:
        blocks.append("Изменения:\n" + "\n\n".join(diffs))
    if merge_request.get("diff_truncated"):
        blocks.append("Примечание: часть изменений усечена из-за размера, полный дифф доступен по ссылке.")

    return "\n\n".join(blocks)


def _gitlab_get(
    base_url: str,
    path: str,
    params: dict | None = None,
    allow_missing: bool = False,
) -> dict | list | None:
    try:
        response = requests.get(
            f"{base_url}/api/v4/{path}",
            headers=_gitlab_headers(),
            params=params,
            timeout=config.REQUEST_TIMEOUT,
        )
    except requests.RequestException as exc:
        raise GitLabError(f"Не удалось подключиться к GitLab: {exc}") from exc
    if response.status_code in (401, 403):
        raise GitLabError("GitLab вернул 401/403. Проверьте GITLAB_TOKEN и доступ к проекту.")
    if response.status_code == 404:
        if allow_missing:
            return None
        raise GitLabError("Merge request не найден. Проверьте ссылку и права токена.")
    if response.status_code >= 400:
        raise GitLabError(f"GitLab API ошибка {response.status_code}")
    return response.json()


def _gitlab_headers() -> dict[str, str]:
    headers = {"Accept": "application/json"}
    if config.GITLAB_TOKEN:
        headers["PRIVATE-TOKEN"] = config.GITLAB_TOKEN
    return headers


def _get_changes(base_url: str, project_id: str, iid: int) -> list[dict]:
    diffs = _gitlab_get(
        base_url,
        f"projects/{project_id}/merge_requests/{iid}/diffs",
        params={"per_page": MAX_DIFF_FILES},
        allow_missing=True,
    )
    if isinstance(diffs, list):
        return diffs
    # GitLab до 15.7: эндпоинт /diffs отсутствует, изменения лежат в /changes
    payload = _gitlab_get(base_url, f"projects/{project_id}/merge_requests/{iid}/changes")
    return (payload or {}).get("changes") or []


def _collect_diffs(changes: list[dict]) -> tuple[list[dict], bool]:
    files: list[dict] = []
    truncated = len(changes) > MAX_DIFF_FILES
    total = 0
    for item in changes[:MAX_DIFF_FILES]:
        diff = (item.get("diff") or "").strip()
        if len(diff) > MAX_FILE_DIFF_CHARS:
            diff = diff[:MAX_FILE_DIFF_CHARS] + "\n... (дифф файла усечен)"
            truncated = True
        if total + len(diff) > MAX_TOTAL_DIFF_CHARS:
            diff = ""
            truncated = True
        total += len(diff)
        files.append(
            {
                "path": item.get("new_path") or item.get("old_path") or "",
                "status": _change_status(item),
                "diff": diff,
            }
        )
    return files, truncated


def _change_status(item: dict) -> str:
    if item.get("new_file"):
        return "добавлен"
    if item.get("deleted_file"):
        return "удален"
    if item.get("renamed_file"):
        return f"переименован из {item.get('old_path', '')}".strip()
    return "изменен"


def extract_text(file_path: str | Path) -> str:
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    raise ParserError(f"Неподдерживаемый формат файла: {ext}")


def _extract_pdf(path: Path) -> str:
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
    if not pages:
        raise ParserError("Файл содержит только изображения, текстовый слой не обнаружен")
    return "\n".join(pages)


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append("\t".join(cells))
        if rows:
            parts.append("\n".join(rows))
    if not parts:
        raise ParserError("Файл не содержит текста")
    return "\n\n".join(parts)
